"""
Sprint Report Generator - Creates comprehensive Word reports for sprints
"""

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List
import logging
from datetime import datetime
from io import BytesIO

logger = logging.getLogger(__name__)


def _convert_to_python_types(obj):
    """Convert numpy types to Python native types."""
    if isinstance(obj, (np.int64, np.int32, np.int16, np.int8)):
        return int(obj)
    elif isinstance(obj, (np.float64, np.float32, np.float16)):
        return float(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif pd.isna(obj):
        return None
    return obj


def set_cell_border(cell, **kwargs):
    """Set cell borders for Word table cells."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:color'), kwargs[edge])
            tcBorders.append(edge_element)
    tcPr.append(tcBorders)


class SprintReportGenerator:
    """Generates comprehensive Word reports for sprints."""
    
    def __init__(self, df: pd.DataFrame, llm_client=None):
        """Initialize with DataFrame and optional LLM client for AI insights."""
        self.df = df.copy()
        self.llm_client = llm_client
        self._preprocess_data()
    
    def _preprocess_data(self):
        """Preprocess data for report generation."""
        date_columns = ['Created_Date', 'Started_Date', 'Completed_Date', 'Sprint_Start', 'Sprint_End']
        for col in date_columns:
            if col in self.df.columns:
                self.df[col] = pd.to_datetime(self.df[col], errors='coerce')
        
        numeric_columns = ['Story_Points', 'Cycle_Time_Days', 'Dev_Time_Hours', 
                          'QA_Time_Hours', 'Estimated_Hours', 'Team_Capacity_Hours']
        for col in numeric_columns:
            if col in self.df.columns:
                self.df[col] = pd.to_numeric(self.df[col], errors='coerce')
    
    def generate_sprint_report(self, sprint_id: str) -> BytesIO:
        """Generate comprehensive Word report for a sprint."""
        sprint_data = self.df[self.df['Sprint_ID'] == sprint_id].copy()
        
        if len(sprint_data) == 0:
            raise ValueError(f"No data found for sprint {sprint_id}")
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Generate all sections
        self._add_cover_page(doc, sprint_data, sprint_id)
        self._add_executive_summary(doc, sprint_data, sprint_id)
        self._add_kpis_table(doc, sprint_data)
        self._add_state_distribution(doc, sprint_data)
        self._add_module_distribution(doc, sprint_data)
        self._add_bugs_deep_dive(doc, sprint_data)
        self._add_cycle_time_analysis(doc, sprint_data)
        self._add_workload_distribution(doc, sprint_data)
        self._add_spillover_analysis(doc, sprint_data)
        self._add_quality_insights(doc, sprint_data)
        self._add_next_sprint_forecast(doc, sprint_data, sprint_id)
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _add_cover_page(self, doc: Document, sprint_data: pd.DataFrame, sprint_id: str):
        """Section 1 - Cover Page."""
        # Title
        title = doc.add_heading(f'Sprint Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(sprint_id, level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Sprint Details
        sprint_start = sprint_data['Sprint_Start'].iloc[0]
        sprint_end = sprint_data['Sprint_End'].iloc[0]
        
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Sprint Period: {sprint_start.strftime('%B %d, %Y')} - {sprint_end.strftime('%B %d, %Y')}\n").bold = True
        
        # Team Members
        team_members = sprint_data['Assignee'].unique()
        doc.add_paragraph()
        team_para = doc.add_paragraph()
        team_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        team_para.add_run('Team Members:\n').bold = True
        team_para.add_run(', '.join(team_members))
        
        # Sprint Goal (derived from modules)
        modules = sprint_data['Area_Module'].unique()
        doc.add_paragraph()
        goal_para = doc.add_paragraph()
        goal_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        goal_para.add_run('Sprint Focus:\n').bold = True
        goal_para.add_run(', '.join(modules))
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, sprint_data: pd.DataFrame, sprint_id: str):
        """Section 2 - Executive Summary."""
        doc.add_heading('Executive Summary', level=1)
        
        # Calculate metrics
        planned_points = sprint_data['Story_Points'].sum()
        completed_points = sprint_data[sprint_data['Status'] == 'Done']['Story_Points'].sum()
        delivery_pct = (completed_points / planned_points * 100) if planned_points > 0 else 0
        
        bugs = sprint_data[sprint_data['Type'] == 'Bug']
        total_bugs = len(bugs)
        fixed_bugs = len(bugs[bugs['Status'] == 'Done'])
        
        spillover_count = len(sprint_data[sprint_data['State'] == 'Spillover'])
        spillover_pct = (spillover_count / len(sprint_data) * 100) if len(sprint_data) > 0 else 0
        
        high_severity_bugs = len(bugs[bugs['Severity'] == 'High']) if 'Severity' in bugs.columns else 0
        
        modules = sprint_data['Area_Module'].value_counts()
        
        # Generate summary text
        summary = doc.add_paragraph()
        summary.add_run(f"The {sprint_id} achieved a delivery rate of {delivery_pct:.1f}%, completing {completed_points:.0f} out of {planned_points:.0f} planned story points. ")
        summary.add_run(f"The team successfully resolved {fixed_bugs} bugs out of {total_bugs} identified, ")
        
        if high_severity_bugs > 0:
            summary.add_run(f"including {high_severity_bugs} high-severity issues. ").font.bold = True
        else:
            summary.add_run("with no high-severity issues remaining. ")
        
        summary.add_run(f"\n\nSpillover: {spillover_count} items ({spillover_pct:.1f}% of sprint scope) were carried over to the next sprint, ")
        if spillover_pct > 15:
            summary.add_run("indicating capacity planning adjustments may be needed. ").font.bold = True
        else:
            summary.add_run("which is within acceptable range. ")
        
        summary.add_run(f"\n\nThe sprint focused primarily on {modules.index[0]} ({modules.values[0]} items)")
        if len(modules) > 1:
            summary.add_run(f" and {modules.index[1]} ({modules.values[1]} items)")
        summary.add_run(". ")
        
        # Key achievements
        doc.add_heading('Key Achievements', level=2)
        achievements = doc.add_paragraph(style='List Bullet')
        achievements.add_run(f"Delivered {completed_points:.0f} story points across {len(sprint_data[sprint_data['Status'] == 'Done'])} items")
        
        if fixed_bugs > 0:
            ach2 = doc.add_paragraph(style='List Bullet')
            ach2.add_run(f"Resolved {fixed_bugs} bugs, improving system stability")
        
        completed_stories = sprint_data[(sprint_data['Type'] == 'Story') & (sprint_data['Status'] == 'Done')]
        if len(completed_stories) > 0:
            ach3 = doc.add_paragraph(style='List Bullet')
            ach3.add_run(f"Completed {len(completed_stories)} user stories enhancing product features")
        
        doc.add_page_break()
    
    def _add_kpis_table(self, doc: Document, sprint_data: pd.DataFrame):
        """Section 3 - Sprint KPIs."""
        doc.add_heading('Sprint KPIs', level=1)
        
        # Calculate KPIs
        planned_points = sprint_data['Story_Points'].sum()
        completed_points = sprint_data[sprint_data['Status'] == 'Done']['Story_Points'].sum()
        delivery_pct = (completed_points / planned_points * 100) if planned_points > 0 else 0
        
        bugs = sprint_data[sprint_data['Type'] == 'Bug']
        bug_severity = bugs['Severity'].value_counts().to_dict() if 'Severity' in bugs.columns else {}
        
        completed = sprint_data[(sprint_data['Status'] == 'Done') & (sprint_data['Cycle_Time_Days'].notna())]
        avg_cycle_time = completed['Cycle_Time_Days'].mean() if len(completed) > 0 else 0
        
        dev_time = sprint_data['Dev_Time_Hours'].sum()
        qa_time = sprint_data['QA_Time_Hours'].sum()
        
        spillover_count = len(sprint_data[sprint_data['State'] == 'Spillover'])
        spillover_pct = (spillover_count / len(sprint_data) * 100) if len(sprint_data) > 0 else 0
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'KPI'
        header_cells[1].text = 'Value'
        header_cells[2].text = 'Explanation'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        kpi_data = [
            ('Planned Story Points', f'{planned_points:.0f}', 'Total story points assigned to sprint'),
            ('Completed Story Points', f'{completed_points:.0f}', 'Story points successfully delivered'),
            ('Sprint Delivery %', f'{delivery_pct:.1f}%', 'Percentage of planned points completed'),
            ('Bug Count', f'{len(bugs)}', 'Total bugs identified in sprint'),
            ('Bug Severity Breakdown', ', '.join([f'{k}: {v}' for k, v in bug_severity.items()]) or 'N/A', 'Distribution by severity'),
            ('Cycle Time Avg', f'{avg_cycle_time:.1f} days', 'Average days from start to completion'),
            ('Dev Time Total', f'{dev_time:.1f} hours', 'Total development hours logged'),
            ('QA Time Total', f'{qa_time:.1f} hours', 'Total QA hours logged'),
            ('Spillover Count', f'{spillover_count}', 'Items carried over to next sprint'),
            ('Spillover %', f'{spillover_pct:.1f}%', 'Percentage of items spilled over'),
        ]
        
        for kpi, value, explanation in kpi_data:
            row_cells = table.add_row().cells
            row_cells[0].text = kpi
            row_cells[1].text = value
            row_cells[2].text = explanation
        
        doc.add_page_break()
    
    def _add_state_distribution(self, doc: Document, sprint_data: pd.DataFrame):
        """Section 4 - State Distribution Analysis."""
        doc.add_heading('State Distribution Analysis', level=1)
        
        state_counts = sprint_data['State'].value_counts()
        total = len(sprint_data)
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'State'
        header_cells[1].text = 'Count'
        header_cells[2].text = '%'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for state, count in state_counts.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(state)
            row_cells[1].text = str(count)
            row_cells[2].text = f'{(count/total*100):.1f}%'
        
        # Analysis
        doc.add_paragraph()
        doc.add_heading('Analysis', level=2)
        
        done_pct = (state_counts.get('Done', 0) / total * 100) if total > 0 else 0
        in_progress_pct = (state_counts.get('In Progress', 0) / total * 100) if total > 0 else 0
        blocked_pct = (state_counts.get('Blocked', 0) / total * 100) if total > 0 else 0
        
        analysis = doc.add_paragraph()
        analysis.add_run(f"The sprint achieved a {done_pct:.1f}% completion rate with {state_counts.get('Done', 0)} items fully delivered. ")
        
        if in_progress_pct > 20:
            analysis.add_run(f"{in_progress_pct:.1f}% of items are still in progress, suggesting scope may have been ambitious. ")
        
        if blocked_pct > 10:
            analysis.add_run(f"Notably, {blocked_pct:.1f}% of items were blocked, indicating dependency management could be improved. ").font.bold = True
        
        if state_counts.get('Spillover', 0) > 0:
            analysis.add_run(f"\n\nSpillover items represent unfinished work that needs to be prioritized in the next sprint. ")
        
        doc.add_page_break()
    
    def _add_module_distribution(self, doc: Document, sprint_data: pd.DataFrame):
        """Section 5 - Module/Area-Wise Distribution."""
        doc.add_heading('Module/Area-Wise Distribution', level=1)
        
        # Group by area
        area_stats = []
        for area in sprint_data['Area_Module'].unique():
            area_data = sprint_data[sprint_data['Area_Module'] == area]
            stories = len(area_data[area_data['Type'] == 'Story'])
            bugs = len(area_data[area_data['Type'] == 'Bug'])
            total_items = len(area_data)
            total_points = area_data['Story_Points'].sum()
            bugs_pct = (bugs / total_items * 100) if total_items > 0 else 0
            
            area_stats.append({
                'area': area,
                'stories': stories,
                'bugs': bugs,
                'total_items': total_items,
                'total_points': total_points,
                'bugs_pct': bugs_pct
            })
        
        # Sort by total points
        area_stats.sort(key=lambda x: x['total_points'], reverse=True)
        
        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        headers = ['Area', 'Stories', 'Bugs', 'Total Items', 'Total Story Points', 'Bugs %']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        for stat in area_stats:
            row_cells = table.add_row().cells
            row_cells[0].text = stat['area']
            row_cells[1].text = str(stat['stories'])
            row_cells[2].text = str(stat['bugs'])
            row_cells[3].text = str(stat['total_items'])
            row_cells[4].text = f"{stat['total_points']:.0f}"
            row_cells[5].text = f"{stat['bugs_pct']:.1f}%"
        
        # Insights
        doc.add_paragraph()
        doc.add_heading('Insights', level=2)
        
        if len(area_stats) > 0:
            insights = doc.add_paragraph()
            insights.add_run(f"• Highest Effort: {area_stats[0]['area']} consumed {area_stats[0]['total_points']:.0f} story points ({area_stats[0]['total_items']} items)\n")
            
            bug_prone = max(area_stats, key=lambda x: x['bugs_pct'])
            if bug_prone['bugs_pct'] > 25:
                insights.add_run(f"• Bug-Prone Module: {bug_prone['area']} has {bug_prone['bugs_pct']:.1f}% bug rate, requiring quality focus\n").bold = True
            
            insights.add_run(f"• The team worked across {len(area_stats)} different modules, indicating broad feature coverage")
        
        doc.add_page_break()
    
    def _add_bugs_deep_dive(self, doc: Document, sprint_data: pd.DataFrame):
        """Section 6 - Bugs Deep-Dive."""
        doc.add_heading('Bugs Deep-Dive', level=1)
        
        bugs = sprint_data[sprint_data['Type'] == 'Bug'].copy()
        
        if len(bugs) == 0:
            doc.add_paragraph("No bugs were reported in this sprint.")
            doc.add_page_break()
            return
        
        # Bugs table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        headers = ['Ticket ID', 'Title', 'Severity', 'Area', 'Status']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        for _, bug in bugs.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(bug['Ticket_ID'])
            row_cells[1].text = str(bug['Title'])[:50]
            row_cells[2].text = str(bug.get('Severity', 'N/A'))
            row_cells[3].text = str(bug['Area_Module'])
            row_cells[4].text = str(bug['Status'])
        
        # Analysis
        doc.add_paragraph()
        doc.add_heading('Bug Analysis', level=2)
        
        severity_dist = bugs['Severity'].value_counts().to_dict() if 'Severity' in bugs.columns else {}
        area_bugs = bugs['Area_Module'].value_counts()
        
        analysis = doc.add_paragraph()
        analysis.add_run(f"Total Bugs: {len(bugs)}\n")
        
        if severity_dist:
            analysis.add_run(f"Severity Distribution: {', '.join([f'{k}: {v}' for k, v in severity_dist.items()])}\n")
        
        if len(area_bugs) > 0:
            analysis.add_run(f"\nMost Bug-Prone Area: {area_bugs.index[0]} ({area_bugs.values[0]} bugs)\n\n")
        
        # Recommendations
        doc.add_heading('Recommendations', level=2)
        recommendations = doc.add_paragraph(style='List Bullet')
        recommendations.add_run("Increase unit test coverage in bug-prone modules")
        
        if severity_dist.get('High', 0) > 2:
            rec2 = doc.add_paragraph(style='List Bullet')
            rec2.add_run("Implement more rigorous code review process for high-severity areas")
        
        rec3 = doc.add_paragraph(style='List Bullet')
        rec3.add_run("Consider automated regression testing to catch issues earlier")
        
        doc.add_page_break()
    
    def _add_cycle_time_analysis(self, doc: Document, sprint_data: pd.DataFrame):
        """Section 7 - Cycle Time Analysis."""
        doc.add_heading('Cycle Time Analysis', level=1)
        
        completed = sprint_data[(sprint_data['Status'] == 'Done') & 
                               (sprint_data['Cycle_Time_Days'].notna())].copy()
        
        if len(completed) == 0:
            doc.add_paragraph("No completed items with cycle time data.")
            doc.add_page_break()
            return
        
        # Cycle time table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        headers = ['Ticket ID', 'Title', 'Story Points', 'Cycle Time (Days)', 'Assignee']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        for _, item in completed.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['Ticket_ID'])
            row_cells[1].text = str(item['Title'])[:40]
            row_cells[2].text = str(int(item['Story_Points']))
            row_cells[3].text = f"{item['Cycle_Time_Days']:.1f}"
            row_cells[4].text = str(item['Assignee'])
        
        # Statistics
        doc.add_paragraph()
        doc.add_heading('Statistics', level=2)
        
        mean_cycle = completed['Cycle_Time_Days'].mean()
        median_cycle = completed['Cycle_Time_Days'].median()
        longest = completed['Cycle_Time_Days'].max()
        shortest = completed['Cycle_Time_Days'].min()
        
        stats = doc.add_paragraph()
        stats.add_run(f"Mean Cycle Time: {mean_cycle:.1f} days\n")
        stats.add_run(f"Median Cycle Time: {median_cycle:.1f} days\n")
        stats.add_run(f"Longest Cycle Time: {longest:.1f} days\n")
        stats.add_run(f"Shortest Cycle Time: {shortest:.1f} days\n")
        
        # Insights
        doc.add_heading('Insights', level=2)
        
        insights = doc.add_paragraph()
        if mean_cycle > 3:
            insights.add_run(f"The average cycle time of {mean_cycle:.1f} days suggests opportunities for process optimization. ")
        
        # Check for bottlenecks
        long_items = completed[completed['Cycle_Time_Days'] > mean_cycle + 2]
        if len(long_items) > 0:
            insights.add_run(f"\n\n{len(long_items)} items took significantly longer than average, potentially indicating scope creep or technical complexity. ")
        
        doc.add_page_break()
    
    def _add_workload_distribution(self, doc: Document, sprint_data: pd.DataFrame):
        """Section 8 - Workload Distribution by Team Member."""
        doc.add_heading('Workload Distribution by Team Member', level=1)
        
        # Calculate workload
        workload = []
        capacity = sprint_data['Team_Capacity_Hours'].iloc[0] if len(sprint_data) > 0 else 160
        
        for assignee in sprint_data['Assignee'].unique():
            assignee_data = sprint_data[sprint_data['Assignee'] == assignee]
            items = len(assignee_data)
            points = assignee_data['Story_Points'].sum()
            dev_hours = assignee_data['Dev_Time_Hours'].sum()
            qa_hours = assignee_data['QA_Time_Hours'].sum()
            total_hours = dev_hours + qa_hours
            capacity_pct = (total_hours / capacity * 100) if capacity > 0 else 0
            
            workload.append({
                'assignee': assignee,
                'items': items,
                'points': points,
                'dev_hours': dev_hours,
                'qa_hours': qa_hours,
                'total_hours': total_hours,
                'capacity_pct': capacity_pct
            })
        
        workload.sort(key=lambda x: x['total_hours'], reverse=True)
        
        # Create table
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        headers = ['Assignee', 'Items', 'Story Points', 'Dev Hours', 'QA Hours', 'Total Hours', '% of Capacity']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        for member in workload:
            row_cells = table.add_row().cells
            row_cells[0].text = member['assignee']
            row_cells[1].text = str(member['items'])
            row_cells[2].text = f"{member['points']:.0f}"
            row_cells[3].text = f"{member['dev_hours']:.1f}"
            row_cells[4].text = f"{member['qa_hours']:.1f}"
            row_cells[5].text = f"{member['total_hours']:.1f}"
            row_cells[6].text = f"{member['capacity_pct']:.1f}%"
        
        # Analysis
        doc.add_paragraph()
        doc.add_heading('Analysis', level=2)
        
        analysis = doc.add_paragraph()
        
        overloaded = [m for m in workload if m['capacity_pct'] > 90]
        underutilized = [m for m in workload if m['capacity_pct'] < 60]
        
        if overloaded:
            analysis.add_run(f"Overloaded Members: {', '.join([m['assignee'] for m in overloaded])} exceeded 90% capacity\n").bold = True
        
        if underutilized:
            analysis.add_run(f"Underutilized Members: {', '.join([m['assignee'] for m in underutilized])} utilized less than 60% capacity\n")
        
        analysis.add_run(f"\nSuggestion: Balance workload distribution in next sprint planning to optimize team productivity.")
        
        doc.add_page_break()
    
    def _add_spillover_analysis(self, doc: Document, sprint_data: pd.DataFrame):
        """Section 9 - Spillover Analysis."""
        doc.add_heading('Spillover Analysis', level=1)
        
        spillover = sprint_data[sprint_data['State'] == 'Spillover'].copy()
        
        if len(spillover) == 0:
            doc.add_paragraph("No items spilled over to the next sprint - excellent execution!")
            doc.add_page_break()
            return
        
        # Spillover table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        headers = ['Ticket ID', 'Title', 'Area', 'Story Points', 'Assignee', 'Status']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        for _, item in spillover.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['Ticket_ID'])
            row_cells[1].text = str(item['Title'])[:40]
            row_cells[2].text = str(item['Area_Module'])
            row_cells[3].text = f"{item['Story_Points']:.0f}"
            row_cells[4].text = str(item['Assignee'])
            row_cells[5].text = str(item['Status'])
        
        # Analysis
        doc.add_paragraph()
        doc.add_heading('Root Cause Analysis', level=2)
        
        analysis = doc.add_paragraph()
        spillover_points = spillover['Story_Points'].sum()
        spillover_pct = (len(spillover) / len(sprint_data) * 100)
        
        analysis.add_run(f"Total Spillover: {len(spillover)} items ({spillover_points:.0f} story points, {spillover_pct:.1f}% of sprint)\n\n")
        
        # By area
        area_spillover = spillover['Area_Module'].value_counts()
        if len(area_spillover) > 0:
            analysis.add_run(f"Most Affected Area: {area_spillover.index[0]} ({area_spillover.values[0]} items)\n")
        
        # By assignee
        assignee_spillover = spillover['Assignee'].value_counts()
        if len(assignee_spillover) > 0:
            analysis.add_run(f"Assignee with Most Spillover: {assignee_spillover.index[0]} ({assignee_spillover.values[0]} items)\n")
        
        # Recommendations
        doc.add_heading('Preventive Recommendations', level=2)
        
        rec = doc.add_paragraph(style='List Bullet')
        rec.add_run("Review sprint capacity planning - consider reducing commitment by 10-15%")
        
        rec2 = doc.add_paragraph(style='List Bullet')
        rec2.add_run("Identify and address blockers earlier in the sprint")
        
        rec3 = doc.add_paragraph(style='List Bullet')
        rec3.add_run("Improve story sizing accuracy through team estimation sessions")
        
        doc.add_page_break()
    
    def _add_quality_insights(self, doc: Document, sprint_data: pd.DataFrame):
        """Section 10 - Quality & Efficiency Insights."""
        doc.add_heading('Quality & Efficiency Insights', level=1)
        
        # Calculate correlations
        completed = sprint_data[(sprint_data['Status'] == 'Done') & 
                               (sprint_data['Cycle_Time_Days'].notna()) &
                               (sprint_data['Story_Points'].notna())]
        
        correlations = doc.add_heading('Correlation Analysis', level=2)
        
        if len(completed) > 1:
            corr_cycle_points = completed['Story_Points'].corr(completed['Cycle_Time_Days'])
            
            corr_para = doc.add_paragraph()
            corr_para.add_run(f"Story Points vs Cycle Time: {corr_cycle_points:.3f}\n")
            
            if corr_cycle_points > 0.7:
                corr_para.add_run("Strong positive correlation - larger stories take proportionally longer (expected)\n")
            elif corr_cycle_points < 0.3:
                corr_para.add_run("Weak correlation - cycle time may be influenced more by complexity than size\n")
        
        # Bug patterns
        bugs = sprint_data[sprint_data['Type'] == 'Bug']
        if len(bugs) > 0:
            bug_area = bugs['Area_Module'].value_counts()
            corr_para.add_run(f"\nBug Concentration: {bug_area.index[0]} has highest bug count, suggesting quality focus needed\n")
        
        # Risk areas
        doc.add_heading('Risk Prediction', level=2)
        
        risk_para = doc.add_paragraph()
        
        spillover_pct = (len(sprint_data[sprint_data['State'] == 'Spillover']) / len(sprint_data) * 100)
        if spillover_pct > 15:
            risk_para.add_run("⚠ High Spillover Risk: Current trend suggests capacity overcommitment\n").bold = True
        
        avg_cycle = completed['Cycle_Time_Days'].mean() if len(completed) > 0 else 0
        if avg_cycle > 4:
            risk_para.add_run("⚠ Cycle Time Risk: Average cycle time exceeds optimal range\n").bold = True
        
        bug_ratio = (len(bugs) / len(sprint_data) * 100)
        if bug_ratio > 20:
            risk_para.add_run("⚠ Quality Risk: Bug ratio exceeds healthy threshold\n").bold = True
        
        # Recommendations
        doc.add_heading('Process Strengthening', level=2)
        
        rec = doc.add_paragraph(style='List Bullet')
        rec.add_run("Implement daily standups focusing on blockers and dependencies")
        
        rec2 = doc.add_paragraph(style='List Bullet')
        rec2.add_run("Introduce mid-sprint health checks to identify at-risk items early")
        
        rec3 = doc.add_paragraph(style='List Bullet')
        rec3.add_run("Enhance automated testing coverage to reduce bug leakage")
        
        doc.add_page_break()
    
    def _add_next_sprint_forecast(self, doc: Document, sprint_data: pd.DataFrame, sprint_id: str):
        """Section 11 - Next Sprint Forecast."""
        doc.add_heading('Next Sprint Forecast', level=1)
        
        # Calculate velocity
        completed_points = sprint_data[sprint_data['Status'] == 'Done']['Story_Points'].sum()
        team_size = sprint_data['Assignee'].nunique()
        
        # Capacity prediction
        doc.add_heading('Expected Delivery Capability', level=2)
        
        forecast = doc.add_paragraph()
        forecast.add_run(f"Based on current sprint performance:\n\n")
        forecast.add_run(f"• Completed Velocity: {completed_points:.0f} story points\n")
        forecast.add_run(f"• Team Size: {team_size} members\n")
        forecast.add_run(f"• Recommended Next Sprint Commitment: {completed_points * 0.9:.0f}-{completed_points:.0f} story points\n\n")
        
        forecast.add_run("The recommendation accounts for a 10% buffer to prevent overcommitment.")
        
        # Risks
        doc.add_heading('Possible Risks', level=2)
        
        risks = doc.add_paragraph(style='List Bullet')
        risks.add_run("Team availability changes (vacation, holidays)")
        
        risks2 = doc.add_paragraph(style='List Bullet')
        
        spillover_count = len(sprint_data[sprint_data['State'] == 'Spillover'])
        if spillover_count > 0:
            risks2.add_run(f"Carrying over {spillover_count} spillover items will reduce new feature capacity")
        else:
            risks2.add_run("External dependencies that could cause delays")
        
        risks3 = doc.add_paragraph(style='List Bullet')
        risks3.add_run("Technical debt accumulation affecting velocity")
        
        # Module hotspots
        doc.add_heading('Module Hotspots', level=2)
        
        module_work = sprint_data.groupby('Area_Module')['Story_Points'].sum().sort_values(ascending=False)
        
        hotspot_para = doc.add_paragraph()
        hotspot_para.add_run("Areas requiring continued focus:\n\n")
        
        for i, (module, points) in enumerate(module_work.head(3).items(), 1):
            hotspot_para.add_run(f"{i}. {module}: {points:.0f} story points\n")
        
        # Test coverage gaps
        doc.add_heading('Test Coverage Gaps', level=2)
        
        bugs = sprint_data[sprint_data['Type'] == 'Bug']
        bug_areas = bugs['Area_Module'].value_counts()
        
        test_para = doc.add_paragraph()
        if len(bug_areas) > 0:
            test_para.add_run(f"Focus testing efforts on: {bug_areas.index[0]}, which had the highest bug concentration.\n\n")
        
        test_para.add_run("Recommendations:\n")
        test_para.add_run("• Increase automated test coverage in bug-prone modules\n")
        test_para.add_run("• Conduct exploratory testing sessions for complex features\n")
        test_para.add_run("• Implement integration testing for cross-module functionality\n")
    
    def get_sprint_list(self) -> List[Dict[str, Any]]:
        """Get list of all sprints with summary information."""
        sprint_list = []
        
        for sprint_id in sorted(self.df['Sprint_ID'].unique()):
            sprint_data = self.df[self.df['Sprint_ID'] == sprint_id]
            
            sprint_start = sprint_data['Sprint_Start'].iloc[0]
            sprint_end = sprint_data['Sprint_End'].iloc[0]
            
            planned_points = sprint_data['Story_Points'].sum()
            completed_points = sprint_data[sprint_data['Status'] == 'Done']['Story_Points'].sum()
            delivery_pct = (completed_points / planned_points * 100) if planned_points > 0 else 0
            
            total_items = len(sprint_data)
            completed_items = len(sprint_data[sprint_data['Status'] == 'Done'])
            
            team_members = sprint_data['Assignee'].unique().tolist()
            
            sprint_list.append({
                'sprint_id': sprint_id,
                'sprint_start': sprint_start.strftime('%Y-%m-%d'),
                'sprint_end': sprint_end.strftime('%Y-%m-%d'),
                'total_items': total_items,
                'completed_items': completed_items,
                'planned_points': float(planned_points),
                'completed_points': float(completed_points),
                'delivery_percentage': round(delivery_pct, 1),
                'team_size': len(team_members),
                'team_members': ', '.join(team_members)
            })
        
        return sprint_list
